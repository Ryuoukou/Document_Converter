from django.shortcuts import render, redirect
from django.core.files.storage import FileSystemStorage
from django.conf import settings
import os
import docx
from docx import Document
import pdfkit
from django.http import HttpResponse  # Импортируем HttpResponse для обработки ошибок
from .models import MyDocument

# Определение функции convert_docx_to_pdf
def convert_docx_to_pdf(docx_path, pdf_path):
    doc = Document(docx_path)
    text = "\n".join([para.text for para in doc.paragraphs])

    temp_txt = 'temp.txt'
    with open(temp_txt, 'w', encoding='utf-8') as f:
        f.write(text)

    pdfkit_config = pdfkit.configuration(
        wkhtmltopdf=r'C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe')

    try:
        pdfkit.from_file(temp_txt, pdf_path, configuration=pdfkit_config)
        if os.path.exists(pdf_path):
            # Конвертация успешно выполнена
            os.remove(temp_txt)  # Удалить временный файл
            success_message = "PDF conversion successful!"
        else:
            # Конвертация не удалась
            return HttpResponse("PDF conversion failed", status=500)
    except Exception as e:
        return HttpResponse(f"Error during PDF conversion: {e}", status=500)


def home(request):
    if request.method == "POST" and request.FILES['document']:
        uploaded_file = request.FILES['document']
        title = request.POST['title']

        fs = FileSystemStorage()
        filename = fs.save(uploaded_file.name, uploaded_file)
        file_path = os.path.join(settings.MEDIA_ROOT, filename)

        if uploaded_file.name.lower().endswith('.docx'):
            pdf_filename = f"{os.path.splitext(filename)[0]}.pdf"
            pdf_full_path = os.path.join(settings.MEDIA_ROOT, 'converted', pdf_filename)

            docx_path = file_path  # Полный путь к .docx файлу
            convert_docx_to_pdf(docx_path, pdf_full_path)  # Вызываем функцию конвертации

            converted_file = pdf_filename
        else:
            converted_file = None

        document = MyDocument(title=title, file=filename, converted_file=converted_file)  # Используем MyDocument
        document.save()

        return redirect('home')

    documents = MyDocument.objects.all()  # Используем MyDocument
    context = {'documents': documents}
    return render(request, 'converter/home.html', context)


